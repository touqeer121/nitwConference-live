import datetime
import math
from django.shortcuts import render, redirect, get_object_or_404
from .models import *
from accounts.models import *
from django.db.models import Q
from django.contrib import messages
import os
import json
import requests
from gdstorage.storage import GoogleDriveStorage, GoogleDrivePermissionType, GoogleDrivePermissionRole, \
	GoogleDriveFilePermission
from django.core.mail import send_mail
from django.conf import settings
import smtplib
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from os.path import basename
from email import encoders
import xlrd
import xlwt
from xlutils.copy import copy
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.models import User
from django.contrib.staticfiles.storage import staticfiles_storage
from nitwConference.settings import STATIC_URL, MEDIA_URL
import stripe
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.db import connection
from docx import Document
from docx.shared import Inches
# import pythoncom
import pathlib
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Table, TableStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from reportlab.lib import colors
from django.core.files import File

permission = GoogleDriveFilePermission(
	GoogleDrivePermissionRole.READER,
	GoogleDrivePermissionType.USER,
	'touqeer.pathan289@gmail.com'
)

def registration_payment(request):
	return render(request, 'registration_payment.html')

@csrf_exempt
def stripe_config(request):
    if request.method == 'GET':
        stripe_config = {'publicKey': settings.STRIPE_PUBLIC_KEY}
        return JsonResponse(stripe_config, safe=False)

@csrf_exempt
def create_checkout_session(request):
    if request.method == 'GET':
        domain_url = 'http://localhost:8000/'
        stripe.api_key = settings.STRIPE_SECRET_KEY
        try:
            print('TRY START')
            # Create new Checkout Session for the order
            # Other optional params include:
            # [billing_address_collection] - to display billing address details on the page
            # [customer] - if you have an existing Stripe Customer ID
            # [payment_intent_data] - capture the payment later
            # [customer_email] - prefill the email input in the form
            # For full details see https://stripe.com/docs/api/checkout/sessions/create

            # ?session_id={CHECKOUT_SESSION_ID} means the redirect will have the session ID set as a query param
            checkout_session = stripe.checkout.Session.create(
                success_url=domain_url + 'success?session_id={CHECKOUT_SESSION_ID}',
                cancel_url=domain_url + 'cancelled/',
                payment_method_types=['alipay'],
                mode='payment',
                line_items=[
                    {
                        'name': 'GCIMB Conference Regular Registration',
                        'quantity': 1,
                        'currency': 'inr',
                        'amount': '55',
                    }
                ]
            )
            print('TRY END')
            return JsonResponse({'sessionId': checkout_session.id})
        except Exception as e:
            print('ECXEPTION')
            return JsonResponse({'error': str(e)})

			
def success(request):
	return render(request, 'registration_success.html')

def cancelled(request):
	return render(request, 'registration_cancelled.html')

def index(request):
	if request.user.is_authenticated: 
		if 'count' in request.session:
			request.session['count'] = request.session['count'] + 1
			print("COUNT = ", request.session['count'])
		else:
			print("NO COUNT")
			
	print ("DATABASE : ",connection.vendor)
	return render(request, 'home.html')

def about_conference(request):
	return render(request, 'about_conference.html')

def organizing_team(request):
	return render(request, 'organizing_team.html')

def who_should_join(request):
	return render(request, 'who_should_join.html')


def committee(request):
	return render(request, 'committee.html')


def track_chairs(request):
	return render(request, 'track_chairs.html')

def preconference_workshop(request):
	return render(request, 'preconference_workshop.html')


def important_dates(request):
	return render(request, 'deadlines.html')


def brochure(request):
	return render(request, 'brochure.html')


def keynote_speakers(request):
	return render(request, 'speakers.html')

def sendReportToAdmin(request, fn , cid, e, gInfo):
	userInfo = 'logged out'
	args = 'no args'
	msg = 'no msg'
	if request.user.is_authenticated : 
		userInfo = request.user.username
	print(e)
	if hasattr(e, 'args'): 
		agr = str(e.args)
	if  hasattr(e, 'message'): 
		msg = str(e.message)	
	else:
		msg = str(e)
	reg = ReceivedException.objects.create(function_name=fn, corresponding_id=cid, exception_args=args, 
			exception_message=msg, general_info=gInfo, current_user_info=userInfo, exception_date=datetime.datetime.now())
	# reg.save()

def forward_registration_info(regID):
	reg = Registration.objects.get(registration_id=regID)
	if reg is None :
		return
	try :
		msg = MIMEMultipart()
		msg.set_unixfrom('author')
		msg['From'] = settings.EMAIL_HOST_USER
		recipients = ['submissions@gcimb.org', 'rama@sgcimb.org', 'ravi@gcimb.org', 'nrustagi@gcimb.org']
		# recipients = ['touqeer.pathan289@gmail.com']
		msg['To'] = ", ".join(recipients)
		msg['Subject'] = 'New participant registered | Registration ID : '+regID

		message = 'Registrations Details : \n' + \
				'Registration ID : '+ str(regID) + '\n' + \
				'Author : '+ str(reg.first_name) + '\n' + \
				'Institute : '+ str(reg.institution) + '\n' + \
				'Address : '+ str(reg.state) +', '+str(reg.country) + '\n' + \
				'Institute : '+ str(reg.institution) + '\n' + \
				'Email : '+ str(reg.email) + '\n' + \
				'Phone : '+ str(reg.phone) + '\n' + \
				'Transaction ID : '+ str(reg.transaction_id) + '\n' + \
				'Registration Date : '+ str(reg.registration_date.date()) + '\n' 

		msg.attach(MIMEText(message))
		mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)
		# mailserver.starttls()
		mailserver.ehlo()
		mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
		mailserver.sendmail(msg['From'], msg['To'], msg.as_string())

		EmailInfo.objects.create(corresponding_id=regID, mail_reason="forward_registration_info", general_info="first", sent_date=datetime.datetime.now())

		print("INFO MAIL SENT\n")
	except Exception as e:
		EmailQueue.objects.create(corresponding_id=str(regID), mail_reason="forward_registration_info", general_info="first", pending_date=datetime.datetime.now())
		sendReportToAdmin(request, "forward_registration_info ", str(regID), e, "")


def is_duplicate_registration(request):
	oldEntry = Registration.objects.filter(first_name=request.POST['fname'], \
				institution= request.POST['institution'])
	# oldEntry=1
	return (len(oldEntry)> 0)

def registration(request):
	if request.method == "POST":
		duplicate = is_duplicate_registration(request)
		abs = None
		if not duplicate:
			regID = None
			try:
				rtype = str(request.POST['registration_type']).strip()
				atype = str(request.POST['author_type']).strip()
				if(atype == '2'):
					atype = '5'
				elif(atype == '3'):
					atype = '6'
				elif(atype == '4'):
					atype = '7'
				elif(atype == '5'):
					atype = '8'
				
				if(rtype != '5' and atype != '8'):
					absID = request.POST['abs_id']
					print("NOT LISTENER OR..",  absID)
					try:
						abs = Abstract.objects.get(abs_id=absID)
					except Abstract.DoesNotExist:
						messages.success(request, "Abstract ID is invalid. Please retry with correct one.")
						print("IF NO ABS ")
						return redirect('/registration')
						
				cnt = Registration_Count.objects.get()
				year = datetime.datetime.now().year
				yy = str(year)
				p1 = yy[2:]
				cnt.registration_count = cnt.registration_count + 1
				p2 = str(cnt.registration_count).zfill(4)
				cnt.save()
				regID = "GCIMBR" + p1 + p2 
				reg = Registration.objects.create(registration_id=regID, registration_date=datetime.datetime.now())
				reg.registration_type = get_object_or_404(Registration_Type, id=rtype)
				reg.author_type = get_object_or_404(Author_Type, id=atype)
				reg.first_name = request.POST['fname']
				if(rtype != '5' and atype != '8'):
					reg.abstract = abs
				
				reg.institution = request.POST['institution']
				reg.country = request.POST['country']
				reg.state = request.POST['state']
				reg.email = request.POST['email']
				reg.phone = request.POST['phone']
				reg.transaction_id = request.POST['trans_id']
				reg.remark = request.POST.get('paymentType', '<ERROR>')
				reg.save()
				messages.success(request, "Thank you for registering. Our team will get back to you in three working days.")
			except Exception as e:
				gInfo = "Error while registration, can be about 'paymentType' value="+str(request.POST['paymentType'])
				sendReportToAdmin(request, "Registration", "Could be : " + str(regID) + " (but is not) ", e, gInfo)
				messages.success(request, "Could not register due to some technical error. Please try again later.")
				return redirect('/registration')
			if regID:
				forward_registration_info(regID)
		else:
			messages.success(request, "You've already registered. Plese wait for our email.")	
		return redirect('/registration')
		
	return render(request, 'register.html')

@csrf_exempt
@login_required(login_url='/sign-in/')
def registration_approval(request):
	context = {}
	if(request.user.username=='gcimb' or request.user.username=='accounts'):
		context['registrations'] = Registration.objects.all().order_by('registration_id')
	return render(request, 'registration_approval.html', context)
	
	messages.success(request, "You do not have the authority to perform this action.")
	return redirect('/')

def update_sheet(absID):
	# response = HttpResponse(content_type='application/ms-excel')
	# response['Content-Disposition'] = 'attachment; filename="Abstracts.xls"'
	rb = xlrd.open_workbook(settings.PROJECT_PATH+'Abstracts.xls',formatting_info=True)
	r_sheet = rb.sheet_by_index(0) 
	r = r_sheet.nrows
	wb = copy(rb) 
	sheet = wb.get_sheet(0) 

	# file = File.objects.all()
	# if file is None:
	# 	file = File.objects.create()
	# else
	# 	file = file[0]
	ppr = get_object_or_404(Abstract, abs_id=absID)
	sheet.write(r,0,absID)
	sheet.write(r,1,ppr.paper_title)
	sheet.write(r,2,ppr.track)
	sheet.write(r,3,ppr.prefix)
	sheet.write(r,4,ppr.first_name)
	sheet.write(r,5,ppr.last_name)
	sheet.write(r,6,ppr.country)
	sheet.write(r,7,ppr.state)
	sheet.write(r,8,ppr.institution)
	sheet.write(r,9,ppr.email)
	sheet.write(r,10,ppr.phone)
	sheet.write(r,11,str(ppr.abstract_pdf.url)[:-16])
	sheet.write(r,12,str(ppr.submission_date).split()[0])
	wb.save('Abstracts.xls')
	print("SHEET UPDATED\n")

def forward_submission_info(absID):
	try:
		ppr = get_object_or_404(Abstract, abs_id=absID)
		msg = MIMEMultipart()
		msg.set_unixfrom('author')
		msg['From'] = settings.EMAIL_HOST_USER
		recipients = ['submissions@gcimb.org', 'rama@gcimb.org', 'ravi@gcimb.org', 'nrustagi@gcimb.org']
		# recipients = ['touqeer.pathan289@gmail.com']
		msg['To'] = ", ".join(recipients)
		msg['Subject'] = 'New Abstract Submitted | Abstract ID : '+absID

		# 'Track : '+ ppr.track + '\n' +\
		message = 'A new abtract has been submitted \n\n' + 'Abstract Details : \n' + \
				'Abstract ID : '+ str(absID) + '\n' + \
				'Author : '+ str(ppr.prefix) + ' ' + str(ppr.first_name) + ' ' + str(ppr.last_name) + '\n' + \
				'Title : '+ str(ppr.paper_title) + '\n' + \
				'Abstract Link : '+ str(ppr.abstract_pdf.url)[:-16] + '\n' + \
				'Address : '+ str(ppr.state) +', '+str(ppr.country) + '\n' + \
				'Institute : '+ str(ppr.institution) + '\n' + \
				'Email : '+ str(ppr.email) + '\n' + \
				'Phone : '+ str(ppr.phone) + '\n' + \
				'Submission Date : '+ str(ppr.submission_date.date()) + '\n' 

		msg.attach(MIMEText(message))
		mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)
		# mailserver.starttls()
		mailserver.ehlo()
		mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
		mailserver.sendmail(msg['From'], msg['To'], msg.as_string())
		
		EmailInfo.objects.create(corresponding_id=absID, mail_reason="forward_submission_info",  general_info="first", sent_date=datetime.datetime.now())

		print("INFO MAIL SENT\n")
	except Exception as e:
		EmailQueue.objects.create(corresponding_id=str(absID), mail_reason="forward_submission_info",  general_info="first", pending_date=datetime.datetime.now())
		sendReportToAdmin(request, "forward_submission_info ", str(absID), e, "")

def forward_paper_submission_info(pprID):
	try : 
		ppr = get_object_or_404(Paper, paper_id=pprID)
		absID = ppr.abstract.abs_id
		msg = MIMEMultipart()
		msg.set_unixfrom('author')
		msg['From'] = settings.EMAIL_HOST_USER
		recipients = ['submissions@gcimb.org', 'rama@gcimb.org', 'ravi@gcimb.org', 'nrustagi@gcimb.org']
		# recipients = ['touqeer.pathan289@gmail.com']
		msg['To'] = ", ".join(recipients)
		msg['Subject'] = 'New Paper Submitted | Paper ID : '+pprID

		message = 'A new paper has been submitted \n\n' + 'Paper Details : \n' + \
				'Paper ID : '+ str(pprID) + '\n' + \
				'Author : '+ str(ppr.prefix) + ' ' + str(ppr.first_name) + ' ' + str(ppr.last_name) + '\n' + \
				'Track : '+ str(ppr.track) + '\n' + \
				'Title : '+ str(ppr.paper_title) + '\n' + \
				'Paper affiliations link : '+ str(ppr.paper_affiliation_pdf.url)[:-16] + '\n' + \
				'Paper manuscript link : '+ str(ppr.paper_manuscript_pdf.url)[:-16] + '\n' + \
				'Address : '+ str(ppr.state) +', '+str(ppr.country) + '\n' + \
				'Institute : '+ str(ppr.institution) + '\n' + \
				'Email : '+ str(ppr.email) + '\n' + \
				'Phone : '+ str(ppr.phone) + '\n' + \
				'Submission Date : '+ str(ppr.submission_date.date()) + '\n' 

		msg.attach(MIMEText(message))
		mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)
		# mailserver.starttls()
		mailserver.ehlo()
		mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
		mailserver.sendmail(msg['From'], msg['To'], msg.as_string())
		
		EmailInfo.objects.create(corresponding_id=absID, mail_reason="forward_paper_submission_info",  general_info="first", sent_date=datetime.datetime.now())

		print("INFO MAIL SENT\n")
	except Exception as e:
		EmailQueue.objects.create(corresponding_id=str(pprID), mail_reason="forward_paper_submission_info",  general_info="first", pending_date=datetime.datetime.now())
		sendReportToAdmin(request, "forward_paper_submission_info ", str(pprID), e, "")
	

def forward_ppt_submission_info(pptID):
	try:
		ppt = get_object_or_404(Ppt, ppt_id=pptID)
		absID = ppt.abstract.abs_id
		msg = MIMEMultipart()
		msg.set_unixfrom('author')
		msg['From'] = settings.EMAIL_HOST_USER
		recipients = ['submissions@gcimb.org', 'rama@gcimb.org', 'ravi@gcimb.org', 'nrustagi@gcimb.org']
		# recipients = ['touqeer.pathan289@gmail.com']
		msg['To'] = ", ".join(recipients)
		msg['Subject'] = 'New Presentation Submitted | Presentation ID : '+pptID

		message = 'A new presentation has been submitted \n\n' + 'Presentation Details : \n' + \
				'Presentation ID : '+ str(pptID) + '\n' + \
				'Author : '+ str(ppt.prefix) + ' ' + str(ppt.first_name) + ' ' + str(ppt.last_name) + '\n' + \
				'Track : '+ str(ppt.track) + '\n' + \
				'Title : '+ str(ppt.ppt_title) + '\n' + \
				'Presentation link : '+ str(ppt.ppt_pdf.url)[:-16] + '\n' + \
				'Address : '+ str(ppt.state) +', '+str(ppt.country) + '\n' + \
				'Institute : '+ str(ppt.institution) + '\n' + \
				'Email : '+ str(ppt.email) + '\n' + \
				'Phone : '+ str(ppt.phone) + '\n' + \
				'Submission Date : '+ str(ppt.submission_date.date()) + '\n' 

		msg.attach(MIMEText(message))
		mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)
		# mailserver.starttls()
		mailserver.ehlo()
		mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
		mailserver.sendmail(msg['From'], msg['To'], msg.as_string())
		
		EmailInfo.objects.create(corresponding_id=absID, mail_reason="forward_ppt_submission_info",  general_info="first", sent_date=datetime.datetime.now())

		print("INFO MAIL SENT\n")
	except Exception as e:
		EmailQueue.objects.create(corresponding_id=str(pptID), mail_reason="forward_ppt_submission_info",  general_info="first", pending_date=datetime.datetime.now())
		sendReportToAdmin(request, "forward_ppt_submission_info ", str(pptID), e, "")	

def is_duplicate_entry(request):
	oldEntry = Abstract.objects.filter(paper_title=request.POST['title'], first_name=request.POST['fname'], last_name=\
	request.POST['lname'], institution= request.POST['institution'])
	return (len(oldEntry)> 0)

def is_duplicate_paper(request):
	try:
		abstract = Abstract.objects.get(abs_id=request.POST['abs_id'])
	except Abstract.DoesNotExist:
		abstract = None
		return false
	oldEntry = Paper.objects.filter(abstract=abstract)
	return (len(oldEntry)> 0)

def is_duplicate_ppt(request):
	try:
		abstract = Abstract.objects.get(abs_id=request.POST['abs_id'])
	except Abstract.DoesNotExist:
		abstract = None
		return false
	oldEntry = Ppt.objects.filter(abstract=abstract)
	return (len(oldEntry)> 0)

def assign_track_chair(request):
	abstracts = Abstract.objects.all()
	for abs in abstracts:
		tracks = TrackChairProfile.objects.filter(track=abs.track)
		if len(tracks) > 0 :
			print("SINGLE TRACK CHAIR FOUND")
			abs.track_A = tracks[0].user.username
		else:
			print("NO TRACK CHAIR FOUND")
		if len(tracks) > 1 :
			print("TWO TRACK CHAIRS FOUND")
			abs.track_B = tracks[1].user.username
		if abs.track == 'Strategic Management and Corporate Governance':
			abs.track_B = 'mahesh'

		abs.save()

def abstract_submission(request):
	if (request.method == "POST"):
		duplicate = is_duplicate_entry(request)
		abs = None
		try:
			if not duplicate:
				doc=request.FILES
				file_pdf = doc['pdf1']
				cnt = Paper_Count.objects.get()
				year = datetime.datetime.now().year
				yy = str(year)
				p1 = yy[2:]
				p2 = str(cnt.paper_count).zfill(4)
				cnt.paper_count = cnt.paper_count + 1
				cnt.save()
				absID = "GCIMB" + p1 + p2
				abs = Abstract.objects.create(abs_id=absID, submission_date=datetime.datetime.now(), abstract_pdf=file_pdf)
				abs.track = request.POST['track']
				abs.prefix = request.POST['prefix']
				abs.first_name = request.POST['fname']
				abs.last_name = request.POST['lname']
				abs.institution = request.POST['institution']
				abs.country = request.POST['country']
				abs.state = request.POST['state']
				abs.email = request.POST['email']
				abs.phone = request.POST['phone']
				abs.paper_title = request.POST['title']
				tracks = TrackChairProfile.objects.filter(track=abs.track)
				if len(tracks) > 0 :
					abs.track_A = tracks[0].user.username
				else:
					print("NO TRACK CHAIR FOUND")
				if len(tracks) > 1 :
					abs.track_B = tracks[1].user.username
				if abs.track == 'Strategic Management and Corporate Governance':
					abs.track_B = 'mahesh'
				abs.save()
			else:
				abs = get_object_or_404(Abstract, paper_title=request.POST['title'])
		except Exception as e:
			sendReportToAdmin(request, "abstract_submission", abs.abs_id +', title => ' + str(request.POST['title']) , e, "exception while creating new abstract, data may not be stored")
			messages.success(request, "Could not submit. Please try again with correct entries.")
			return redirect('/abstract-submission')
		
		try:
			msg = MIMEMultipart()
			msg.set_unixfrom('author')
			msg['From'] = settings.EMAIL_HOST_USER
			msg['To'] = request.POST['email']

			if not duplicate:
				msg['Subject'] = 'Abstract submission acknowledgement'
				message = 'Hello ' + str(abs.prefix) + ' ' + str(abs.first_name) + ' ' + str(abs.last_name) + ',\n\n' + \
						'Hope you are safe and doing well. This is to acknowledge that we have received your abstract.' + \
						'Your abstract ID will be ' + str(abs.abs_id) +'. Please make a note of it and quote the same in future communications.\n\n' + \
						'Your abstract will be sent for review and you should be hearing from us very soon on the next steps.\n\n' + \
						'Many thanks for considering to submit your work to GCIMB.\n\n'+\
						'Best Regards,\n' + \
						'Organizing Team,\n' + \
						'Global Conference on Innovations in Management and Business'
			else:
				msg['Subject'] = 'Abstract already submitted'
				#abs cannot be None here
				message = 'Looks like your abstract with the title \"'+ str(abs.title) +'\" (Abstract ID : ' + str(abs.abs_id) + ') is already ' + \
					'submitted and you should receive an email. \nIn case you didn’t, please write to submissions@gcimb.org quoting your details.'	
			
			msg.attach(MIMEText(message))

			mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)
			# mailserver.starttls()
			mailserver.ehlo()
			mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
			mailserver.sendmail(msg['From'], msg['To'], msg.as_string())
	
			if not duplicate:
				EmailInfo.objects.create(corresponding_id=abs.abs_id, mail_reason="abstract_submission",  general_info="first", sent_date=datetime.datetime.now())
				forward_submission_info(abs.abs_id)
				messages.success(request, "You've successfully submitted the abstract.")
			else:
				EmailInfo.objects.create(corresponding_id=abs.abs_id, mail_reason="abstract_submission",  general_info="duplicate", sent_date=datetime.datetime.now())
				messages.success(request, "You've already submitted an abstract with this title.")
			print("EVRYTHING DONE\n")
			
		except Exception as e:
			if not duplicate:
				EmailQueue.objects.create(corresponding_id=abs.abs_id, mail_reason="abstract_submission",  general_info="first", pending_date=datetime.datetime.now())
				messages.success(request, "You've successfully submitted the abstract.")
			else:
				EmailQueue.objects.create(corresponding_id=abs.abs_id, mail_reason="abstract_submission",  general_info="duplicate", pending_date=datetime.datetime.now())
				messages.success(request, "You've already submitted an abstract with this title.")
			sendReportToAdmin(request, "abstract_submission", str(abs.abs_id)+', title => ' + str(abs.paper_title), e, "exception while mailing, data is stored")
		return redirect('/abstract-submission')
	return render(request, 'abstract_submission.html')


def paper_submission(request):
	if (request.method == "POST"):
		duplicate = is_duplicate_paper(request)
		ppr = None
		abs = None
		try:
			if not duplicate:
				doc=request.FILES
				file_pdf1 = doc['pdf1']
				file_pdf2 = doc['pdf2']
				absID = request.POST['abs_id']
				try:
					abs = Abstract.objects.get(abs_id=absID)
				except Abstract.DoesNotExist:
					messages.success(request, "Abstract ID is invalid. Please retry with correct one.")
					return redirect('/paper-submission')
				cnt = Full_Paper_Count.objects.get()
				year = datetime.datetime.now().year
				yy = str(year)[2:]
				p1 = yy
				cnt.full_paper_count = cnt.full_paper_count + 1
				p2 = str(cnt.full_paper_count).zfill(4)
				cnt.save()
				pprID = "GCIMBP" + p1 + p2
				ppr = Paper.objects.create(paper_id=pprID, submission_date=datetime.datetime.now(), paper_affiliation_pdf=file_pdf1, \
						paper_manuscript_pdf=file_pdf2, abstract = abs)
				ppr.track = request.POST['track']
				ppr.prefix = request.POST['prefix']
				ppr.first_name = request.POST['fname']
				ppr.last_name = request.POST['lname']
				ppr.institution = request.POST['institution']
				ppr.country = request.POST['country']
				ppr.state = request.POST['state']
				ppr.email = request.POST['email']
				ppr.phone = request.POST['phone']
				ppr.paper_title = request.POST['title']

				tracks = TrackChairProfile.objects.filter(track=ppr.track)
				if len(tracks) > 0 :
					ppr.track_A = tracks[0].user.username
				else:
					print("NO TRACK CHAIR FOUND")
				if len(tracks) > 1 :
					ppr.track_B = tracks[1].user.username
				if ppr.track == 'Strategic Management and Corporate Governance':
					ppr.track_B = 'mahesh'
				ppr.save()
			else:
				ppr = get_object_or_404(Paper, abstract = abs)	
		except Exception as e:
			sendReportToAdmin(request, "paper_submission", str(ppr.paper_id)+', title => ' + str(ppr.paper_title), e, "exception while creating new paper, data may not be stored")
			messages.success(request, "Could not submit. Please try agin with correct entries.")
			return redirect('/paper-submission')
		
		try:
			msg = MIMEMultipart()
			msg.set_unixfrom('author')
			msg['From'] = settings.EMAIL_HOST_USER
			msg['To'] = request.POST['email']

			if not duplicate:
				msg['Subject'] = 'Paper submission acknowledgement'
				message = 'Hello ' + ppr.prefix + ' ' + ppr.first_name + ' ' + ppr.last_name + ',\n\n' + \
						'Hope you are safe and doing well. This is to acknowledge that we have received your paper.' + \
						'Your paper ID will be ' + str(ppr.paper_id) +'. Please make a note of it and quote the same in future communications.\n\n' + \
						'Many thanks for considering to submit your work to GCIMB.\n\n'+\
						'Best Regards,\n' + \
						'Organizing Team,\n' + \
						'Global Conference on Innovations in Management and Business'
			else:
				msg['Subject'] = 'Paper already submitted'
				#ppr cannot be None
				message = 'Looks like your paper with the title \"'+ ppr.paper_title +'\" (Paper ID : ' + ppr.paper_id + ') is already ' + \
					'submitted and you should receive an email. \nIn case you didn’t, please write to submissions@gcimb.org quoting your details.'	
			
			msg.attach(MIMEText(message))
			mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)

			mailserver.ehlo()
			mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
			mailserver.sendmail(msg['From'], msg['To'], msg.as_string())

			if not duplicate:
				EmailInfo.objects.create(corresponding_id=str(ppr.paper_id), mail_reason="paper_submission",  general_info="first", sent_date=datetime.datetime.now())
				forward_paper_submission_info(str(ppr.paper_id))
				messages.success(request, "You've successfully submitted the paper.")
			else: 
				EmailInfo.objects.create(corresponding_id=str(ppr.paper_id), mail_reason="paper_submission",  general_info="duplicate", sent_date=datetime.datetime.now())
				messages.success(request, "You've already submitted a paper for this abstract.")
			print("EVRYTHING DONE with PAPER SUBMISSION\n")

		except Exception as e:
			if not duplicate:
				EmailQueue.objects.create(corresponding_id=str(ppr.paper_id), mail_reason="paper_submission",  general_info="first", pending_date=datetime.datetime.now())
				messages.success(request, "You've successfully submitted the paper.")
			else: 
				EmailQueue.objects.create(corresponding_id=str(ppr.paper_id), mail_reason="paper_submission",  general_info="duplicate", pending_date=datetime.datetime.now())
				messages.success(request, "You've already submitted a paper for this abstract.")
			sendReportToAdmin(request, "paper_submission", str(ppr.paper_id)+', title => ' + str(ppr.paper_title), e, "exception while mailing, data is stored")

		return redirect('/paper-submission')
	return render(request, 'paper_submission.html')	

def ppt_submission(request):
	if (request.method == "POST"):
		duplicate = is_duplicate_ppt(request)
		ppt = None
		try:
			if not duplicate:
				doc=request.FILES
				file_pdf = doc['pdf']
				absID = request.POST['abs_id']
				try:
					abs = Abstract.objects.get(abs_id=absID)
				except Abstract.DoesNotExist:
					messages.success(request, "Abstract ID is invalid. Please retry with correct one.")
					return redirect('/ppt-submission')
				cnt = Ppt_Count.objects.get()
				year = datetime.datetime.now().year
				yy = str(year)[2:]
				p1 = yy
				cnt.ppt_count = cnt.ppt_count + 1
				p2 = str(cnt.ppt_count).zfill(4)
				cnt.save()
				pptID = "GCIMBT" + p1 + p2
				ppt = Ppt.objects.create(ppt_id=pptID, submission_date=datetime.datetime.now(), ppt_pdf=file_pdf, abstract = abs)
				ppt.track = request.POST['track']
				ppt.prefix = request.POST['prefix']
				ppt.first_name = request.POST['fname']
				ppt.last_name = request.POST['lname']
				ppt.institution = request.POST['institution']
				ppt.country = request.POST['country']
				ppt.state = request.POST['state']
				ppt.email = request.POST['email']
				ppt.phone = request.POST['phone']
				ppt.ppt_title = request.POST['title']

				tracks = TrackChairProfile.objects.filter(track=ppt.track)
				if len(tracks) > 0 :
					ppt.track_A = tracks[0].user.username
				else:
					print("NO TRACK CHAIR FOUND")
				if len(tracks) > 1 :
					ppt.track_B = tracks[1].user.username
				if ppt.track == 'Strategic Management and Corporate Governance':
					ppt.track_B = 'mahesh'
				ppt.save()
			else:
				ppt = get_object_or_404(Ppt, abstract = abs)
		except:
			sendReportToAdmin(request, "ppt_submission", str(ppt.ppt_id) +', title => ' + str(ppt.ppt_title), e, "exception while creating new ppt, data may not be stored")
			messages.success(request, "Could not submit. Please try again with correct entries.")
			return redirect('/ppt-submission')
		
		try:
			msg = MIMEMultipart()
			msg.set_unixfrom('author')
			msg['From'] = settings.EMAIL_HOST_USER
			msg['To'] = request.POST['email']

			if not duplicate:
				msg['Subject'] = 'Presentation submission acknowledgement'
				message = 'Hello ' + ppt.prefix + ' ' + ppt.first_name + ' ' + ppt.last_name + ',\n\n' + \
						'Hope you are safe and doing well. This is to acknowledge that we have received your presentation.' + \
						'Your paper ID will be ' + ppt.ppt_id +'. Please make a note of it and quote the same in future communications.\n\n' + \
						'Many thanks for considering to submit your work to GCIMB.\n\n'+\
						'Best Regards,\n' + \
						'Organizing Team,\n' + \
						'Global Conference on Innovations in Management and Business'
			else:
				msg['Subject'] = 'Presentation already submitted'
				message = 'Looks like your presentation with the title \"'+ ppt.ppt_title +'\" (Presentation ID : ' + ppt.ppt_id + ') is already ' + \
					'submitted and you should receive an email. \nIn case you didn’t, please write to submissions@gcimb.org quoting your details.'	
			
			msg.attach(MIMEText(message))
			mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)

			mailserver.ehlo()
			mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
			mailserver.sendmail(msg['From'], msg['To'], msg.as_string())


			if not duplicate:
				EmailInfo.objects.create(corresponding_id=ppt.ppt_id, mail_reason="ppt_submission",  general_info="first", sent_date=datetime.datetime.now())
				forward_ppt_submission_info(ppt.ppt_id)
				messages.success(request, "You've successfully submitted the ppt.")
			else:
				EmailInfo.objects.create(corresponding_id=ppt.ppt_id, mail_reason="ppt_submission",  general_info="duplicate", sent_date=datetime.datetime.now())
				messages.success(request, "You've already submitted a ppt for this abstract.")
			print("EVRYTHING DONE with PPT SUBMISSION\n")

		except Exception as e:
			if not duplicate:
				EmailQueue.objects.create(corresponding_id=ppt.ppt_id, mail_reason="ppt_submission",  general_info="first", pending_date=datetime.datetime.now())
				messages.success(request, "You've successfully submitted the ppt.")
			else:
				EmailQueue.objects.create(corresponding_id=ppt.ppt_id, mail_reason="ppt_submission",  general_info="duplicate", pending_date=datetime.datetime.now())
				messages.success(request, "You've already submitted a ppt for this abstract.")
			sendReportToAdmin(request, "ppt_submission", str(ppt.ppt_id)+', title => ' + str(ppt.ppt_title), e, "exception while mailing, data is stored")

		return redirect('/ppt-submission')
	return render(request, 'ppt_submission.html')	


def early_bird(request):
	return render(request, 'early-bird.html')


def non_early(request):
	return render(request, 'non-early.html')


def call_for_papers(request):
	return render(request, 'call_for_papers.html')


def abstract_submission_guidelines(request):
	return render(request, 'abstract_submission_guidelines.html')


def publication_opportunities(request):
	return render(request, 'publication_opportunities.html')


def evaluation_process(request):
	return render(request, 'evaluation_process.html')


def contact_us(request):
	if request.method == 'POST':
		try:
			msg = MIMEMultipart()
			msg.set_unixfrom('author')
			msg['From'] = 'info@gcimb.org'
			msg['To'] = request.POST['email']
			msg['Subject'] = request.POST['subject']
			message = request.POST['message']
			msg.attach(MIMEText(message))
			mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)
			mailserver.ehlo()
			mailserver.login('info@gcimb.org', 'info123@gcimb')
			mailserver.sendmail(msg['From'], msg['To'], msg.as_string())
			mailserver.quit()
			EmailInfo.objects.create(str(request.POST['email']), mail_reason="contact_us", general_info=str(request.POST['subject']) +'\n'+ str(request.POST['message']), sent_date=datetime.datetime.now())
			messages.success(request, "Message Sent successfully. We'll get back to you soon.")
		except Exception as e:
			sendReportToAdmin(request, "contact_us", 'Subject : '+ str(request.POST['subject']), e, str(request.POST['message']))
		return redirect('/contact_us')
	return render(request, 'contact_us.html')

def digital_transformation_and_information_systems(request):
	return render(request, 'digital_transformation_and_information_systems.html')

@login_required(login_url='/sign-in/')
def export_abstracts_sheet(request):
	assign_track_chair(request)
	response = HttpResponse(content_type='application/ms-excel')
	response['Content-Disposition'] = 'attachment; filename="Abstracts.xls"'
	wb = xlwt.Workbook(encoding='utf-8')
	ws = wb.add_sheet('Abstracts')
	
	# Sheet header, first row
	row_num = 0
	
	font_style = xlwt.XFStyle()
	font_style.font.bold = True
	columns = ['Abstract ID', 'Title', 'Track', 'Prefix', 'First Name', 'Last Name', 'Country', 'State', \
			'Institute', 'Email', 'Phone', 'File Name', 'File Link', 'Submission Date']		
	for col_num in range(len(columns)):
		ws.write(row_num, col_num, columns[col_num], font_style)
	
	# Sheet body, remaining rows
	font_style = xlwt.XFStyle()
	
	rows = Abstract.objects.all().values_list('abs_id', 'paper_title', 'track', 'prefix', 'first_name', 'last_name', 'country', 'state',\
			 'institution', 'email', 'phone', 'abstract_pdf', 'submission_date' )

	counter = 0
	for row in rows:
		row_num += 1
		for col_num in range(len(row)):
			if col_num==13:
				dt = str(row[col_num]).split()[0]
				ws.write(row_num, col_num, dt, font_style)
			elif col_num==11:
				ppr = get_object_or_404(Abstract, abs_id=row[0])
				if ppr:
					ws.write(row_num, col_num, str(ppr.abstract_pdf), font_style)
			elif col_num==12:
				ppr = get_object_or_404(Abstract, abs_id=row[0])
				if ppr:
					ws.write(row_num, col_num, str(ppr.abstract_pdf.url)[:-16], font_style)
			else:
				ws.write(row_num, col_num, row[col_num], font_style)
		counter += 1
	wb.save(response)
	return response

@csrf_exempt
@login_required(login_url='/sign-in/')
def remark_abstracts(request):
	print("START")
	context = {}
	if(request.user.username=='gcimb'):
		print("SUPERADMIN")
		context['abstracts'] = Abstract.objects.all().order_by('abs_id')
		context['track'] = ''
		return render(request, 'all_abstracts.html', context)
	if(request.user.username=='shekar'):
		print("SHEKAR")
		context['abstracts'] = Abstract.objects.filter(track='Digital Transformation and Information Systems')
		return render(request, 'all_abstracts.html', context)

	curUser = request.user
	curUser = TrackChairProfile.objects.get(user=curUser)
	print("USER " , curUser)

	
	if(request.user.username=='mahesh'):
		track2 = 'Strategic Management and Corporate Governance'
	else:
		track2 = 'NA' 
	context['abstracts'] = Abstract.objects.filter(Q(track=curUser.track) | Q(track=track2)).order_by('abs_id') 
	context['track'] = curUser.track
	context['track2'] = track2
	return render(request, 'all_abstracts.html', context)
	
@login_required(login_url='/sign-in/')
def approve_abstract(request, abstractid):
	cd = get_object_or_404(Abstract, pk=abstractid)
	uname = request.user.username
	remark = request.POST.get('remark')
	if uname=='gcimb':
		if cd.is_finally_approved:
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract approved.")
			cd.is_finally_rejected = False
			cd.is_finally_approved = True
		cd.remark = str(remark).strip()
		cd.save()
	elif uname=='shekar':
		if cd.status_C == '1':
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract approved.")
			cd.status_C = '1'
		cd.remark_C = str(remark).strip()
		cd.save()
	elif uname == cd.track_A:
		if cd.is_approved_by_A:
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract approved.")
			cd.is_rejected_by_A = False
			cd.is_approved_by_A = True
		cd.remark_A = str(remark).strip()
		cd.save()
	elif uname == cd.track_B:
		if cd.is_approved_by_B:
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract approved.")
			cd.is_rejected_by_B = False
			cd.is_approved_by_B = True
		cd.remark_B = str(remark).strip()
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	print("A | Remark : ", remark)
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def reject_abstract(request, abstractid):
	cd = get_object_or_404(Abstract, pk=abstractid)
	uname = request.user.username
	remark = request.POST.get('remark')
	if uname=='gcimb':
		if cd.is_finally_rejected:
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract rejected.")
			cd.is_finally_rejected = True
			cd.is_finally_approved = False
		cd.remark = str(remark).strip()
		cd.save()
	elif uname=='shekar':
		if cd.status_C == '0':
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract rejected.")
			cd.status_C = '0'
		cd.remark_C = str(remark).strip()
		cd.save()
	elif uname == cd.track_A:
		if cd.is_rejected_by_A:
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract rejected.")
			cd.is_rejected_by_A = True
			cd.is_approved_by_A = False
		cd.remark_A = str(remark).strip()
		cd.save()
	elif uname == cd.track_B:
		if cd.is_rejected_by_B:
			messages.success(request, "Remarks Updated.")
		else:
			messages.success(request, "Abstract rejected.")
			cd.is_rejected_by_B = True
			cd.is_approved_by_B = False
		cd.remark_B = str(remark).strip()
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	print("R | Remark : ", remark)
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def remove_remark(request, abstractid):
	cd = get_object_or_404(Abstract, pk=abstractid)
	uname = request.user.username
	remark = request.POST.get('remark')
	if uname=='gcimb':
		if not cd.is_finally_approved and not cd.is_finally_rejected:
			messages.success(request, "Already no remarks.")
		else:
			messages.success(request, "Remarks Removed.")
			cd.is_finally_approved = cd.is_finally_rejected = False
		cd.remark = ''
		cd.save()
	elif uname=='shekar':
		if cd.status_C == '2':
			messages.success(request, "Already no remarks.")
		else:
			messages.success(request, "Remarks Removed.")
			cd.status_C = '2'
		cd.remark_C = ''
		cd.save()
	elif uname == cd.track_A:
		if not cd.is_approved_by_A and not cd.is_rejected_by_A:
			messages.success(request, "Already no remarks.")
		else:
			messages.success(request, "Remarks Removed.")
			cd.is_approved_by_A = cd.is_rejected_by_A = False
		cd.remark_A = ''
		cd.save()
	elif uname == cd.track_B:
		if not cd.is_approved_by_B and not cd.is_rejected_by_B:
			messages.success(request, "Already no remarks.")
		else:
			messages.success(request, "Remarks Removed.")
			cd.is_approved_by_B = cd.is_rejected_by_B = False
		cd.remark_B = ''
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	print("R | Remark : ", remark)
	return redirect(request.META.get('HTTP_REFERER', '/'))


@login_required(login_url='/sign-in/')
def approve_id(request, registrationid):
	cd = get_object_or_404(Registration, pk=registrationid)
	uname = request.user.username
	# remark = request.POST.get('remark')
	if uname=='gcimb' or uname =='accounts':
		if cd.id_status == '1':
			messages.success(request, "ID already approved.")
		else:
			messages.success(request, "ID approved.")
			cd.id_status = '1'
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def reject_id(request, registrationid):
	cd = get_object_or_404(Registration, pk=registrationid)
	uname = request.user.username
	# remark = request.POST.get('remark')
	if uname=='gcimb' or uname =='accounts':
		if cd.id_status == '0':
			messages.success(request, "ID already rejected.")
		else:
			messages.success(request, "ID rejected.")
			cd.id_status = '0'
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def reset_decision_for_id(request, registrationid):
	cd = get_object_or_404(Registration, pk=registrationid)
	uname = request.user.username
	# remark = request.POST.get('remark')
	if uname=='gcimb' or uname =='accounts':
		messages.success(request, "Decision reset.")
		cd.id_status = '2'
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def generate_receipt(request, registrationid):
	try:
		reg = get_object_or_404(Registration, pk=registrationid)
	except ECXEPTION as e:
		messages.success("Error while receipt generation! regisration could not be found in database.")
		return

	invoice = "GCIMBIN21" + str(reg.registration_id)[-4:]
	filePath = 'static/files/' + reg.registration_id 
	fileName = reg.registration_id + '_receipt.pdf'
	pathlib.Path(filePath).mkdir(exist_ok=True) 

	filePath += '/'
	title = 'Receipt'
	logoPath = 'static/images/gcimb_img.png' 
	signPath = 'static/images/sign.jpeg'

	date = 'Date: ' + str(reg.registration_date.date())		

	regards0 = 'Convenor'
	regards = 'Finance Team'
	regards2 = 'GCIMB 2021'

	width, height = landscape(letter)
	styles = getSampleStyleSheet()
	styleN = styles["Normal"]
	styleN.alignment = TA_LEFT
	styleN.fontName = 'Courier'
	styleN.fontSize = 16
	styleN.leading = 16
	
	styleS = styles["Code"]
	styleS.alignment = TA_CENTER
	styleS.fontName = 'Courier'
	styleS.fontSize = 16
	styleS.textColor = 'white'

	styleBH = styles["BodyText"]
	styleBH.alignment = TA_CENTER
	styleBH.fontName = 'Courier'
	styleBH.fontSize = 16

	def coord(x, y, unit=1):
		x, y = x * unit, height -  y * unit
		return x, y

	
	invoiceNo = Paragraph('''<b>Invoice Number:</b>''', styleN)
	invoiceNoVal = Paragraph(f'''{invoice}''', styleN)
	gap = Paragraph(f'''''', styleN)
	name = Paragraph('''<b>Participant\'s Name : </b>''', styleN)
	nameVal = Paragraph(f'''{reg.first_name}''', styleN)
	institution = Paragraph('''<b>Institution / Organization:</b>''', styleN)
	institutionVal = Paragraph(f'''{reg.institution}''', styleN)
	mode = Paragraph('''<b>Mode of payment:</b>''', styleN)
	modeVal = Paragraph(f'''Account transfer''', styleN)
	gap = Paragraph(f'''''', styleN)


	# Headers
	hdescrpcion = Paragraph('''<b>Item</b>''', styleS)
	hpartida = Paragraph('''<b>Quantity</b>''', styleS)
	hcandidad = Paragraph('''<b>Amount</b>''', styleS)

	# Texts
	descrpcion = Paragraph(f'''{reg.registration_type.registration_type}''', styleBH)
	qua = Paragraph('1', styleBH)
	amount = Paragraph(f'''{reg.remark}''', styleBH)

	data= [
		[invoiceNo, invoiceNoVal],
			[gap, gap],
		[name, nameVal],
			[gap, gap],
		[institution, institutionVal],
			[gap, gap],
		[mode, modeVal],
			[gap, gap],
	]

	mainTable = Table(data, colWidths=[8 * cm, 16 * cm])

	mainTable.setStyle(TableStyle([
							('VALIGN', (0,0), (-1,-1), 'MIDDLE')
						# ('INNERGRID', (0,0), (-1,-1), 0.25, colors.black),
						# ('BOX', (0,0), (-1,-1), 0.25, colors.black),
						]))

	data= [[hdescrpcion, hpartida, hcandidad],
		[descrpcion, qua, amount]
	]
	table = Table(data, colWidths=[8 * cm, 8 * cm, 8 * cm])

	table.setStyle(TableStyle([
						('INNERGRID', (0,0), (-1,-1), 0.25, colors.grey),
						('BOX', (0,0), (-1,-1), 0.25, colors.grey),
						('BACKGROUND', (0,0), (-1,0), colors.Color(red=(67.0/255),green=(202.0/255),blue=(232.0/255))),
						('BOTTOMPADDING', (0,0), (-1,-1), 12)
						]))


	pdf = canvas.Canvas(filePath + fileName, pagesize=landscape(letter))
	pdf.rect(20, 20, 750, 570, stroke=1, fill=0) 
	pdf.rect(22, 22, 746, 566, stroke=1, fill=0) 

	pdf.setTitle(title)
	
	pdf.setFont('Courier', 36)
	pdf.drawString(50, 550, title)

	pdf.setFont('Courier', 16)
	pdf.drawString(550, 500, date)

	pdf.drawInlineImage(logoPath, 550, 534)

	pdf.line(30, 530, 760, 530)

	pdf.setFont('Courier', 16)
	pdf.drawString(550, 65, regards0)
	pdf.drawString(550, 50, regards)
	pdf.drawString(550, 35, regards2)
	
	pdf.drawInlineImage(signPath, 550, 80)
	
	mainTable.wrapOn(pdf, width-20, height-20)
	mainTable.drawOn(pdf, *coord(1.8, 11.5, cm))


	table.wrapOn(pdf, width-20, height-20)
	table.drawOn(pdf, *coord(1.8, 15, cm))

	pdf.save()

	local_file = open(filePath +'/'+ fileName, "rb")
	django_file = File(local_file)
	
	#delete old receipts
	Receipt.objects.filter(registration=reg).delete()		

	#create new
	receipt = Receipt.objects.create(registration=reg, receipt_file=django_file)
	receipt.save()

	return

@login_required(login_url='/sign-in/')
def approve_payment(request, registrationid):
	cd = get_object_or_404(Registration, pk=registrationid)
	uname = request.user.username
	# remark = request.POST.get('remark')
	if uname=='gcimb' or uname =='accounts':
		if cd.payment_status == '1':
			messages.success(request, "Payment already approved, Receipt regenerated")
			# try:
			# 	fileName = 'static/files/' + cd.registration_id + '/' + cd.registration_id + '_receipt.pdf'
			# 	os.remove(fileName)
			# 	print('old recipt file deleted')
			# except OSError:
			# 	print('old recipt file NOT deleted')
			# 	pass
			generate_receipt(request, registrationid)
			print("IF me")
		else:
			messages.success(request, "Payment approved, Receipt generated.")
			generate_receipt(request, registrationid)
			print("ELSE me")
			cd.payment_status = '1'
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def reject_payment(request, registrationid):
	cd = get_object_or_404(Registration, pk=registrationid)
	uname = request.user.username
	# remark = request.POST.get('remark')
	if uname=='gcimb' or uname =='accounts':
		if cd.payment_status == '0':
			messages.success(request, "Payment already rejected.")
		else:
			messages.success(request, "Payment rejected.")
			cd.payment_status = '0'
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def reset_decision_for_payment(request, registrationid):
	cd = get_object_or_404(Registration, pk=registrationid)
	uname = request.user.username
	# remark = request.POST.get('remark')
	if uname=='gcimb' or uname =='accounts':
		messages.success(request, "Decision reset.")
		cd.payment_status = '2'
		cd.save()
	else : 
		messages.success(request, "You do not have the authority to perform this action.")
	
	return redirect(request.META.get('HTTP_REFERER', '/'))


@login_required(login_url='/sign-in/')
def test_doc_old(request):
	uname = request.user.username
	if uname=='gcimb':
		reg = get_object_or_404(Registration, pk='GCIMBR210003')
		document = Document()

		section = document.sections[0]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = new_width
		section.page_height = new_height

		sec_pr = section._sectPr # get the section properties el
		# create new borders el
		pg_borders = OxmlElement('w:pgBorders')
		# specifies how the relative positioning of the borders should be calculated
		pg_borders.set(qn('w:offsetFrom'), 'page')
		for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
			border_el = OxmlElement(f'w:{border_name}')
			border_el.set(qn('w:val'), 'double') # a single line
			border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
			border_el.set(qn('w:space'), '24')
			border_el.set(qn('w:color'), 'auto')
			pg_borders.append(border_el) # register single border to border el
		sec_pr.append(pg_borders) # apply border changes to section

		header = document.sections[0].header
		paragraph = header.paragraphs[0]

		text_run = paragraph.add_run()
		text_run.text = "Receipt" + '\t\t                                      '  # For center align of text
		text_run.bold = True
		font = text_run.font
		font.name = 'Arial'
		font.size = Pt(20)

		logo_run = paragraph.add_run()
		logo_run.add_picture('static/images/'+'gcimb_img.png', width=Inches(2))

		
		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		date = p.add_run('\nDate: '+ str(reg.registration_date.date())+'\n')
		# date.font.name = 'Arial'
		date.font.size = Pt(14)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.LEFT
		name = p.add_run('Participant’s Name: ')
		name.bold = True
		val = p.add_run(str(reg.first_name))
		# name.font.name = 'Arial'
		name.font.size = val.font.size = Pt(14)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.LEFT
		name = p.add_run('Institution / Organization: ')
		name.bold = True
		val = p.add_run(str(reg.institution))
		# name.font.name = 'Arial'
		name.font.size = val.font.size = Pt(14)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.LEFT
		name = p.add_run('Mode of payment: ')
		name.bold = True
		val = p.add_run('Account transfer')
		# name.font.name = 'Arial'
		name.font.size = val.font.size = Pt(14)

		print("PATH : " + STATIC_URL + 'images/' + 'sign_tmp.png')

		table = document.add_table(rows=1, cols=3)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Item'
		hdr_cells[1].text = 'Quantity'
		hdr_cells[2].text = 'Amount'
		
		row_cells = table.add_row().cells
		row_cells[0].text = reg.registration_type.registration_type
		row_cells[1].text = '1'
		row_cells[2].text = reg.remark
		
		for row in table.rows:
			for cell in row.cells:
				paragraphs = cell.paragraphs
				for paragraph in paragraphs:
					for run in paragraph.runs:
						font = run.font
						font.size= Pt(14)

		document.add_paragraph('\n')

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		sign = p.add_run()
		# sign.add_picture('static/images/'+'sign_tmp.png', width=Inches(1.25))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		name = p.add_run('For Finance Team')
		name.font.name = 'Arial'
		name.font.size = Pt(14) 

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		gcimb = p.add_run('GCIMB 2021')
		gcimb.font.name = 'Arial'
		gcimb.font.size = Pt(15) 

		# document.add_page_break()
		document.save("static/files/"+'demo.docx')
		# pythoncom.CoInitialize()
		# convert("media/demo.docx", "media/demoC.pdf")
		messages.success(request, "Doc saved successfully")

	else : 
		messages.success(request, "You do not have the authority to perform this action.")

	return redirect(request.META.get('HTTP_REFERER', '/'))


@login_required(login_url='/sign-in/')
def send_approval_mail(request, registrationid):
	try: 
		reg = get_object_or_404(Registration, pk=registrationid)
		msg = MIMEMultipart()
		msg.set_unixfrom('author')
		msg['From'] = settings.EMAIL_HOST_USER
		msg['To'] = reg.email.strip()
		# msg['To'] = 'touqeer.pathan4567@gmail.com'

		msg['Subject'] = 'Registration Aprroved!'
		message = 'Hello ' + reg.first_name + ',\n\n' + \
				'Hope you are safe and doing well. This is to acknowledge that your registration has been approved. Please find the attached receipt and keep it for future reference.\n\n' + \
				'Many thanks for considering to attend the conference.\n\n'+\
				'Best Regards,\n' + \
				'Organizing Team,\n' + \
				'Global Conference on Innovations in Management and Business\n' + \
				'Receipt link: ' + str(reg.Receipt.receipt_file.url)[:-16]
		
		msg.attach(MIMEText(message))

		# filePath =  'static/files/'+ reg.registration_id + '/' + reg.registration_id + '_receipt.pdf'
		# with open(filePath, "rb") as fil:
		# 	part = MIMEApplication(
		# 		fil.read(),
		# 		Name=basename(filePath)
		# 	)
		# After the file is closed
		# part['Content-Disposition'] = 'attachment; filename="%s"' % basename(filePath)
		# msg.attach(part)

		# msg.attach_file('abc.pdf', static('files/'+ reg.registration_id + '/' + reg.registration_id + '_receipt.pdf'))
		mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)

		mailserver.ehlo()
		mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
		mailserver.sendmail(msg['From'], msg['To'], msg.as_string())
		
		messages.success(request, 'Email Sent Successfully')
		EmailInfo.objects.create(corresponding_id=registrationid, mail_reason="send_approval_mail", general_info="", sent_date=datetime.datetime.now())
	except Exception as e:
		messages.success(request, 'Email could not be sent. (Limit Exceeded) OR ' + str(e.message))
		EmailQueue.objects.create(corresponding_id=registrationid, mail_reason="send_approval_mail", general_info="", pending_date=datetime.datetime.now())
		sendReportToAdmin(request, "send_approval_mail" , reg.registration_id, e, "")

	return redirect(request.META.get('HTTP_REFERER', '/'))
	

@login_required(login_url='/sign-in/')
def test_mail(request):
	reg = get_object_or_404(Registration, pk='GCIMBR210003')
	msg = MIMEMultipart()
	msg.set_unixfrom('author')
	msg['From'] = settings.EMAIL_HOST_USER
	msg['To'] = 'touqeer.pathan289@gmail.com'

	msg['Subject'] = 'Registration Aprroved!'
	message = 'Hello ' + reg.first_name + ',\n\n' + \
			'Hope you are safe and doing well. This is to acknowledge that we have received your presentation.' + \
			'Your registration has been approved. Please find the attached receipt and keep it for future reference.\n\n' + \
			'Many thanks for considering to attend the conference.\n\n'+\
			'Best Regards,\n' + \
			'Organizing Team,\n' + \
			'Global Conference on Innovations in Management and Business'	
	
	msg.attach(MIMEText(message))

	filePath =  'static/files/'+ reg.registration_id + '/' + reg.registration_id + '_receipt.pdf'
	with open(filePath, "rb") as fil:
		part = MIMEApplication(
			fil.read(),
			Name=basename(filePath)
		)
	# After the file is closed
	part['Content-Disposition'] = 'attachment; filename="%s"' % basename(filePath)
	msg.attach(part)

	# msg.attach_file('abc.pdf', static('files/'+ reg.registration_id + '/' + reg.registration_id + '_receipt.pdf'))
	mailserver = smtplib.SMTP_SSL('smtpout.secureserver.net', 465)

	mailserver.ehlo()
	mailserver.login(settings.EMAIL_HOST_USER, settings.EMAIL_HOST_PASSWORD)
	mailserver.sendmail(msg['From'], msg['To'], msg.as_string())

	return redirect(request.META.get('HTTP_REFERER', '/'))

@login_required(login_url='/sign-in/')
def test_doc(request):
	uname = request.user.username
	if uname=='gcimb':
		reg = get_object_or_404(Registration, pk='GCIMBR210003')

		filePath = 'static/files/' + reg.registration_id 
		fileName = reg.registration_id + '_receipt.pdf'
		pathlib.Path(filePath).mkdir(exist_ok=True) 

		filePath += '/'
		title = 'Receipt'
		logoPath = 'static/images/gcimb_img.png' 
		signPath = 'static/images/sign.jpeg'

		date = 'Date: ' + str(reg.registration_date.date())		

		regards0 = 'Convenor'
		regards = 'Finance Team'
		regards2 = 'GCIMB 2021'

		width, height = landscape(letter)
		styles = getSampleStyleSheet()
		styleN = styles["Normal"]
		styleN.alignment = TA_LEFT
		styleN.fontName = 'Courier'
		styleN.fontSize = 16
		styleN.leading = 16
		
		styleS = styles["Code"]
		styleS.alignment = TA_CENTER
		styleS.fontName = 'Courier'
		styleS.fontSize = 16
		styleS.textColor = 'white'

		styleBH = styles["BodyText"]
		styleBH.alignment = TA_CENTER
		styleBH.fontName = 'Courier'
		styleBH.fontSize = 16

		def coord(x, y, unit=1):
			x, y = x * unit, height -  y * unit
			return x, y

		name = Paragraph('''<b>Participant\'s Name : </b>''', styleN)
		nameVal = Paragraph(f'''{reg.first_name}''', styleN)
		institution = Paragraph('''<b>Institution / Organization:</b>''', styleN)
		institutionVal = Paragraph(f'''{reg.institution}''', styleN)
		mode = Paragraph('''<b>Mode of payment:</b>''', styleN)
		modeVal = Paragraph(f'''Account transfer''', styleN)
		gap = Paragraph(f'''''', styleN)
		mode = Paragraph('''<b>Invoice Number:</b>''', styleN)
		modeVal = Paragraph(f'''GCIMBIN210001''', styleN)
		gap = Paragraph(f'''''', styleN)


		# Headers
		hdescrpcion = Paragraph('''<b>Item</b>''', styleS)
		hpartida = Paragraph('''<b>Quantity</b>''', styleS)
		hcandidad = Paragraph('''<b>Amount</b>''', styleS)
	
		# Texts
		descrpcion = Paragraph(f'''{reg.registration_type.registration_type}''', styleBH)
		qua = Paragraph('1', styleBH)
		amount = Paragraph(f'''{reg.remark}''', styleBH)

		data= [[name, nameVal],
				[gap, gap],
			[institution, institutionVal],
				[gap, gap],
			[mode, modeVal],
				[gap, gap],
		]

		mainTable = Table(data, colWidths=[8 * cm, 16 * cm])

		mainTable.setStyle(TableStyle([
								('VALIGN', (0,0), (-1,-1), 'MIDDLE')
							# ('INNERGRID', (0,0), (-1,-1), 0.25, colors.black),
							# ('BOX', (0,0), (-1,-1), 0.25, colors.black),
							]))

		data= [[hdescrpcion, hpartida, hcandidad],
			[descrpcion, qua, amount]
		]
		table = Table(data, colWidths=[8 * cm, 8 * cm, 8 * cm])

		table.setStyle(TableStyle([
							('INNERGRID', (0,0), (-1,-1), 0.25, colors.grey),
							('BOX', (0,0), (-1,-1), 0.25, colors.grey),
							('BACKGROUND', (0,0), (-1,0), colors.Color(red=(67.0/255),green=(202.0/255),blue=(232.0/255))),
							('BOTTOMPADDING', (0,0), (-1,-1), 12)
							]))


		pdf = canvas.Canvas(filePath + fileName, pagesize=landscape(letter))
		pdf.rect(20, 20, 750, 570, stroke=1, fill=0) 
		pdf.rect(22, 22, 746, 566, stroke=1, fill=0) 

		pdf.setTitle(title)
		
		pdf.setFont('Courier', 36)
		pdf.drawString(50, 550, title)

		pdf.setFont('Courier', 16)
		pdf.drawString(550, 500, date)

		pdf.drawInlineImage(logoPath, 550, 534)

		pdf.line(30, 530, 760, 530)

		pdf.setFont('Courier', 16)
		pdf.drawString(550, 65, regards0)
		pdf.drawString(550, 50, regards)
		pdf.drawString(550, 35, regards2)
		
		pdf.drawInlineImage(signPath, 550, 80)
		
		mainTable.wrapOn(pdf, width-20, height-20)
		mainTable.drawOn(pdf, *coord(1.8, 11, cm))


		table.wrapOn(pdf, width-20, height-20)
		table.drawOn(pdf, *coord(1.8, 15, cm))

		pdf.save()
		messages.success(request, 'file saved!')
	else:
		messages.success(request, 'Access Denied!')
	return redirect('/')

@login_required(login_url='/sign-in/')
def test_doc3(request):
	uname = request.user.username
	if uname=='gcimb':
		reg = get_object_or_404(Registration, pk='GCIMBR210003')

		filePath = 'static/files/' + reg.registration_id 
		fileName = reg.registration_id + '_receipt.pdf'
		
		width, height = landscape(letter)
		styles = getSampleStyleSheet()
		styleN = styles["BodyText"]
		styleN.alignment = TA_LEFT
		
		styleBH = styles["Normal"]
		styleBH.alignment = TA_CENTER


		def coord(x, y, unit=1):
			x, y = x * unit, height -  y * unit
			return x, y

		# Headers
		hdescrpcion = Paragraph('''<b>Item</b>''', styleBH)
		hpartida = Paragraph('''<b>Quantity</b>''', styleBH)
		hcandidad = Paragraph('''<b>Amount</b>''', styleBH)

		# Texts
		descrpcion = Paragraph(f'''{reg.registration_type.registration_type}''', styleBH)
		qua = Paragraph('1', styleBH)
		amount = Paragraph(f'''{reg.remark}''', styleBH)

		data= [[hdescrpcion, hcandidad,hcandidad],
			[descrpcion, qua, amount]
		]

		table = Table(data, colWidths=[8 * cm, 8 * cm, 8 * cm])

		table.setStyle(TableStyle([
							('INNERGRID', (0,0), (-1,-1), 0.25, colors.black),
							('BOX', (0,0), (-1,-1), 0.25, colors.black),
							]))

		c = canvas.Canvas(filePath +'/'+ fileName, pagesize=landscape(letter))
		table.wrapOn(c, width, height)
		table.drawOn(c, *coord(1.8, 9.6, cm))
		c.save()
		messages.success(request, 'file saved!')
	else:
		messages.success(request, 'Access Denied!')
	return redirect('/')
